"""Update the client user guide with the v2 formula breakdown.

The client asked for the full weighting per category written into the guide so he can
explain the system to his own users. Every number here is READ FROM THE CODE at build
time (propline.scoring.WEIGHTS and the module constants) rather than typed in, so the
document cannot quietly drift away from what the pipeline actually does.

Run after changing any weight or threshold:

    python scripts/build_guide.py

Idempotent: section 7 is torn down and rebuilt on every run, so re-running after a
weight change updates the tables instead of duplicating them.

Requires python-docx. There is no pandoc or LibreOffice on this machine, so the .docx is
edited in place rather than regenerated from markdown — which also preserves the v1
document's styling and the sections that are still accurate.
"""

from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

sys.path.insert(0, str(Path(__file__).resolve().parents[1]))

from propline import profiles as P  # noqa: E402
from propline import scoring as S  # noqa: E402
from propline import weather as W  # noqa: E402

GUIDE = Path("PropLine MLB - User Guide.docx")

# Plain-English meaning for every signal, keyed exactly as in scoring.WEIGHTS. The
# weight itself is never written here — it is read from the code.
MEANINGS = {
    "hits": {
        "recent_xwoba": ("Recent expected on-base quality",
                         "What his contact has actually been worth lately, judged on how "
                         "hard and at what angle he hits the ball — not just whether it "
                         "happened to fall in."),
        "contact_rate": ("Contact rate",
                         "How often he makes contact when he swings. A hit needs contact "
                         "first."),
        "ld_sweet_index": ("Line-drive and sweet-spot rate",
                           "How often he hits the ball at the angles that actually "
                           "produce hits."),
        "starter_whip": ("Opposing starter's WHIP",
                         "How many baserunners that starter allows per inning."),
    },
    "total_bases_power": {
        "iso_recent_14day": ("Isolated power, last 14 days",
                             "Recent extra-base power (slugging minus batting average)."),
        "recent_barrel_pct": ("Barrel rate, recent",
                              "Share of batted balls struck at the ideal combination of "
                              "speed and angle."),
        "iso_season": ("Isolated power, season",
                       "The same power measure across the full season, as a stabiliser."),
        "starter_slg_allowed_vs_hand": ("Starter's slugging allowed to this side",
                                        "How hard left- or right-handed hitters "
                                        "specifically have hit this starter."),
        "recent_hard_hit": ("Hard-hit rate, recent",
                            "Share of batted balls hit at 95 mph or more."),
    },
    "total_bases_volume": {
        "recent_hit_rate": ("Recent hit rate", "How often he has been getting hits."),
        "contact_rate": ("Contact rate", "How often he makes contact when he swings."),
        "matchup_est_woba": ("Matchup expected wOBA",
                             "His record against the specific pitches this starter "
                             "throws, weighted by how often each is thrown."),
        "lineup_spot": ("Batting-order spot",
                        "Higher in the order means more plate appearances."),
    },
    "home_runs": {
        "hr_matchup_rv": ("Arsenal matchup run value",
                          "His run value against each pitch type, weighted by how often "
                          "this starter actually throws it to his side of the plate."),
        "recent_barrel_pct": ("Barrel rate, recent",
                              "The best single available indicator of home-run power."),
        "fly_ball_rate": ("Fly-ball rate",
                          "A home run has to be in the air. Ground-ball hitters are "
                          "discounted here."),
        "iso_recent_14day": ("Isolated power, last 14 days",
                             "Recent extra-base power."),
        "recent_hard_hit": ("Hard-hit rate, recent",
                            "Share of batted balls at 95 mph or more."),
    },
    "rbis": {
        "matchup_est_slg": ("Matchup expected slugging",
                            "How much damage he projects to do against this starter's "
                            "particular mix of pitches."),
        "recent_tb_rate": ("Recent total-base rate", "Recent extra-base production."),
        "rbi_spot": ("Batting-order spot",
                     "The 3-4-5 hitters come to bat with men on base most often."),
        "season_est_slg": ("Season expected slugging",
                           "The full-season version, as a stabiliser."),
    },
    "runs": {
        "matchup_est_woba": ("Matchup expected on-base quality",
                             "You cannot score a run without first reaching base."),
        "run_spot": ("Batting-order spot",
                     "The 1-2-3 hitters score most often — the heart of the order bats "
                     "behind them."),
        "recent_hit_rate": ("Recent hit rate", "How often he has been reaching lately."),
        "matchup_k_pct": ("Strikeout risk in this matchup",
                          "Counts AGAINST the score: a strikeout cannot score a run."),
    },
    "strikeouts": {
        "split_k_matchup": ("Strikeout rate vs tonight's lineup",
                            "His strikeout rate against left- and right-handed hitters, "
                            "weighted by who is actually in the lineup he faces."),
        "whiff_14day": ("Whiff rate, last 14 days",
                        "How often hitters swing and miss against him right now."),
        "opp_lineup_k_pct": ("How strikeout-prone that lineup is",
                             "Some lineups strike out far more often than others."),
        "whip_efficiency": ("WHIP efficiency",
                            "Whether he works efficiently enough to stay in the game "
                            "long enough to pile up strikeouts."),
    },
    "team_total": {
        "lineup_matchup_woba": ("This lineup vs the opposing starter",
                                "Every hitter's matchup against the man they face, "
                                "averaged across the lineup."),
        "opp_starter_weak": ("How hittable that starter has been",
                             "What he gives up generally, beyond this specific matchup."),
        "opp_pen_workload": ("Opposing bullpen workload, last 3 days",
                             "Pitch counts and innings for the relief unit. A worked "
                             "bullpen leaks runs from the sixth inning on."),
        "park_runs": ("Park run factor", "100 is neutral; higher favours scoring."),
        "recent_team_form": ("Recent team form",
                             "How the lineup as a whole has been producing."),
        "opp_starter_whip": ("Opposing starter's WHIP",
                             "More baserunners allowed means more scoring chances."),
    },
    "game_total": {
        "combined_offense": ("Both lineups combined",
                             "The two teams' matchup quality added together."),
        "combined_starter_weak": ("Both starters' hittability",
                                  "How much the two starting pitchers give up."),
        "combined_pen_workload": ("Both bullpens' workload, last 3 days",
                                  "Two tired bullpens is the classic route to a total "
                                  "going over late."),
        "park_runs": ("Park run factor", "100 is neutral; higher favours scoring."),
        "combined_starter_k9": ("Combined starter strikeout rate",
                                "Counts AGAINST the score. Two strikeout pitchers "
                                "suppress a total, because strikeouts remove balls in "
                                "play entirely. This is the pitching-duel check."),
    },
}

TABLE_ORDER = [
    ("hits", "Hits"),
    ("total_bases_power", "Total Bases — power path"),
    ("total_bases_volume", "Total Bases — volume path"),
    ("home_runs", "Home Runs"),
    ("rbis", "RBIs"),
    ("runs", "Runs"),
    ("strikeouts", "Strikeouts (pitchers)"),
    ("team_total", "Team Totals"),
    ("game_total", "Game Totals"),
]


def _set_borders(table):
    """Light grid lines. The v1 tables are borderless, but these carry numbers."""
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:color"), "BFBFBF")
        borders.append(el)
    table._tbl.tblPr.append(borders)


def main() -> int:
    if not GUIDE.exists():
        print(f"missing {GUIDE}")
        return 1

    doc = Document(str(GUIDE))
    body = doc.element.body

    def find(*prefixes, style=None):
        """First paragraph starting with any of `prefixes`.

        Several of the edits below rewrite the very text they searched for, so the
        anchors have to survive their own rewrite — hence the alternatives. Without
        this the script worked once and then failed on every re-run, which defeats the
        point of generating the weights from code.
        """
        for p in doc.paragraphs:
            if (style is None or p.style.name == style) and any(
                    p.text.strip().startswith(x) for x in prefixes):
                return p
        raise LookupError(" / ".join(prefixes))

    def retext(p, text):
        """Replace text, keeping the paragraph's style and its first run's format."""
        for run in list(p.runs)[1:]:
            run._r.getparent().remove(run._r)
        if p.runs:
            p.runs[0].text = text
        else:
            p.add_run(text)

    # ---------------------------------------------------------------- section 6
    v2_bullets = {
        "Hits —": "Hits — recent expected on-base quality, contact rate, line-drive and "
                  "sweet-spot contact, and the opposing starter's WHIP. Ground-ball "
                  "hitters are penalised.",
        "Total Bases —": "Total Bases — scored twice, once for power and once for volume, "
                         "and the better of the two is used. A hitter can be a strong "
                         "play either way.",
        "Home Runs —": "Home Runs — a pitch-by-pitch matchup matrix, barrel rate, "
                       "fly-ball rate, recent power and hard-hit rate. Capped when the "
                       "starter strongly suppresses hard contact.",
        "RBIs —": "RBIs — matchup slugging, recent total bases, batting-order spot and "
                  "season slugging, discounted when the hitters ahead of him rarely "
                  "reach base.",
        "Runs —": "Runs — matchup on-base quality, batting-order spot and recent hit "
                  "rate, minus strikeout risk, with the same lineup-context discount.",
        "Strikeouts —": "Strikeouts — his strikeout rate against this lineup's mix of "
                        "left- and right-handed hitters, 14-day whiff rate, how "
                        "strikeout-prone the lineup is, and WHIP efficiency, minus a "
                        "penalty for an expected short outing.",
        "Game / Team Totals —": "Game / Team Totals — both offences against the opposing "
                                "starters, three-day bullpen workload, the ballpark, a "
                                "pitching-duel check and first-pitch temperature.",
    }
    for prefix, new in v2_bullets.items():
        retext(find(prefix), new)

    # Bullpen is a real workload model in v2, not the v1 availability check. Anchored on
    # the (stable) Heading 2 rather than on the body text this rewrites.
    from docx.text.paragraph import Paragraph
    pen_head = find("Bullpen", style="Heading 2")
    pen_paras, el = [], pen_head._p.getnext()
    while el is not None and len(pen_paras) < 2 and el.tag.endswith("}p"):
        p = Paragraph(el, pen_head._parent)
        if p.style.name.startswith("Heading"):
            break
        pen_paras.append(p)
        el = el.getnext()

    if len(pen_paras) >= 1:
        retext(pen_paras[0],
               "The system adds up how many pitches and how many innings each bullpen "
               "has thrown over the last three days and labels the unit Rested, Average "
               f"or Overworked. Under {P.PEN_RESTED_MAX_PITCHES} pitches across three "
               f"days is Rested; over {P.PEN_TIRED_MIN_PITCHES} is Overworked.")
    if len(pen_paras) >= 2:
        retext(pen_paras[1],
               "Those thresholds were measured from a full 30-team slate rather than "
               "picked by hand. An overworked bullpen pushes game and team totals up, "
               "because the innings from the sixth on are the ones most likely to leak "
               "runs. The board shows the status for each side.")

    # ------------------------------------------------- section 7: rebuilt wholesale
    h7 = find("7.", style="Heading 1")
    h8 = find("8.")

    # Tear down everything between the two headings so a re-run replaces rather than
    # duplicates. This is what makes the script safe to run after any weight change.
    el = h7._p.getnext()
    while el is not None and el is not h8._p:
        nxt_el = el.getnext()
        body.remove(el)
        el = nxt_el

    cursor = h7._p

    def place(el):
        nonlocal cursor
        cursor.addnext(el)
        cursor = el
        return el

    def para(text, style=None, bold=False):
        p = doc.add_paragraph(style=style)
        p.add_run(text).bold = bold
        return place(p._p)

    para("Every score is 0–100 and is relative to the other players on that day's slate, "
         "not an absolute rating. For each signal every player is ranked against everyone "
         "else playing that day, and those rankings are then combined using the weights "
         "below. A score of 80 means the play sits near the top of today's board — it is "
         "not an 80% chance of anything.")
    para("The weights in each table add up to 1.00. A negative weight means the signal "
         "counts against the score rather than for it.")

    for key, title in TABLE_ORDER:
        para(title, style="Heading 2")
        weights = S.WEIGHTS[key]
        rows = sorted(weights.items(), key=lambda kv: -abs(kv[1]))

        table = doc.add_table(rows=len(rows) + 1, cols=3)
        _set_borders(table)
        for i, head in enumerate(("Signal", "Weight", "What it means")):
            c = table.rows[0].cells[i]
            c.text = head
            c.paragraphs[0].runs[0].bold = True
        for r, (col, w) in enumerate(rows, start=1):
            label, meaning = MEANINGS[key].get(col, (col, ""))
            table.rows[r].cells[0].text = label
            table.rows[r].cells[1].text = f"{w:.2f}" if w > 0 else f"−{abs(w):.2f}"
            table.rows[r].cells[2].text = meaning
        place(table._tbl)

    para("Adjustments that sit outside the weights", style="Heading 2")
    para("A few things are not weighted signals but caps, penalties or multipliers "
         "applied after the weighted score is worked out:")
    for text in (
        f"Ground-ball penalty (Hits) — above a {P.GB_PENALTY_THRESHOLD:.0f}% ground-ball "
        f"rate, up to {P.GB_PENALTY_MAX:.0f} points come off, scaled by how far over the "
        "line he is. Balls on the ground become outs and double plays.",

        "Two paths (Total Bases) — the power and volume scores are worked out separately "
        "and the higher one is used. The board shows which path a pick came from. A "
        f"power score above {S.TB_STRICT_THRESHOLD:.0f} also flags the pick as a "
        "stronger candidate for two or more bases.",

        f"Home-run cap — a home-run score is capped at {S.HR_SUPPRESSION_CAP:.0f} when "
        "the opposing starter strongly suppresses hard contact, however good the hitter "
        "looks on his own numbers.",

        f"Lineup context (Runs and RBIs) — the on-base ability of the "
        f"{P.TABLE_SETTERS} hitters batting ahead of him applies a multiplier between "
        f"{P.CONTEXT_MIN:.2f} and {P.CONTEXT_MAX:.2f}. It is a discount only: weak "
        "table-setters pull a score down, but strong ones never inflate it.",

        f"Short-leash penalty (Strikeouts) — if a starter is expected to throw fewer "
        f"than {S.SHORT_LEASH_PITCHES} pitches, up to {S.LEASH_PENALTY_MAX:.0f} points "
        "come off. Strikeouts need innings.",

        f"Weather (Totals) — first-pitch temperature adjusts the total by at most "
        f"{(W.TEMP_MULT_MAX - 1) * 100:.0f}% either way, neutral at "
        f"{W.TEMP_NEUTRAL_F:.0f}°F. Warm air is thinner and the ball carries. Indoor "
        "parks are always neutral.",
    ):
        para(text, style="List Paragraph")

    para("If you think a weight is wrong, tell me the category and what you would move. "
         "Changing them is a small edit, and this section is generated from the code, so "
         "it updates automatically when they change.")

    print(f"  ok    section 7 rebuilt: {len(TABLE_ORDER)} weight tables")

    # ------------------------------------------------------- corrections elsewhere
    # v1 shipped "8." as body text, which is why the contents jumped 7 to 9.
    h8.style = doc.styles["Heading 1"]

    # Weather is no longer deferred, so that bullet is replaced by the real limitation.
    retext(find("Weather —", "Wind —"),
           "Wind — shown, but deliberately not scored. Whether wind helps or hurts "
           "depends on which way a ballpark faces, and that orientation is not published "
           "in any source this system can read. Temperature is scored; wind speed and "
           "direction are displayed so you can apply your own judgement on parks you "
           "know.")

    # The tool does now read the market.
    retext(find("These are rankings, not predictions"),
           "These are rankings, not predictions. The tool tells you which plays the "
           "numbers favour today. It also pulls the sportsbook's posted game totals and "
           "strikeout lines where they are available and shows where the model and the "
           "market disagree — but it still cannot tell you what to bet.")

    retext(find("Nothing here is locked"),
           "Nothing here is locked. The weights, the number of picks shown, the "
           "categories and the run times are all adjustable.")

    print("  ok    corrections applied (heading 8, wind, market, changes)")
    doc.save(str(GUIDE))
    print(f"  ok    {GUIDE}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
